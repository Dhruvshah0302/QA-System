import streamlit as st
import os
import logging
from pathlib import Path
from dotenv import load_dotenv
import warnings

# Suppress deprecation warnings for cleaner UI
warnings.filterwarnings('ignore', category=DeprecationWarning)

# Set up basic logging
logging.basicConfig(level=logging.INFO)

# Load .env file from current directory
load_dotenv()

# Try to import LlamaIndex
try:
    from llama_index.core import SimpleDirectoryReader, VectorStoreIndex, Settings, Document
    from llama_index.llms.gemini import Gemini
    from llama_index.embeddings.gemini import GeminiEmbedding
    from llama_index.core.node_parser import SentenceSplitter
    import google.generativeai as genai
    LLAMA_INDEX_AVAILABLE = True
except ImportError as e:
    LLAMA_INDEX_AVAILABLE = False
    st.error(f"Import error: {str(e)}")

# Document processing
from pypdf import PdfReader
import docx

def extract_text_from_pdf(file):
    """Extract text from PDF file"""
    try:
        pdf_reader = PdfReader(file)
        text = ""
        for page in pdf_reader.pages:
            text += page.extract_text() + "\n"
        return text
    except Exception as e:
        st.error(f"Error reading PDF: {str(e)}")
        return ""

def extract_text_from_docx(file):
    """Extract text from DOCX file"""
    try:
        doc = docx.Document(file)
        text = ""
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"
        return text
    except Exception as e:
        st.error(f"Error reading DOCX: {str(e)}")
        return ""

def extract_text_from_txt(file):
    """Extract text from TXT file"""
    try:
        return file.read().decode('utf-8')
    except Exception as e:
        st.error(f"Error reading TXT: {str(e)}")
        return ""

def load_data_from_upload(uploaded_file):
    """Load data from uploaded file"""
    try:
        logging.info("Data loading started...")
        
        if uploaded_file is None:
            return ""
        
        file_extension = uploaded_file.name.split('.')[-1].lower()
        
        if file_extension == 'pdf':
            text = extract_text_from_pdf(uploaded_file)
        elif file_extension == 'docx':
            text = extract_text_from_docx(uploaded_file)
        elif file_extension == 'txt':
            text = extract_text_from_txt(uploaded_file)
        else:
            st.error("Unsupported file format")
            return ""
        
        logging.info("Data loading completed...")
        return text
        
    except Exception as e:
        logging.error(f"Exception in loading data: {str(e)}")
        st.error(f"Error loading data: {str(e)}")
        return ""

def create_index_from_text(text, api_key):
    """Create LlamaIndex from text using Gemini"""
    try:
        if not LLAMA_INDEX_AVAILABLE:
            st.error("LlamaIndex is not available")
            return None
            
        # Configure Gemini
        genai.configure(api_key=api_key)
        
        # Set up Settings with Gemini - exactly like your notebook
        Settings.llm = Gemini(model="gemini-2.5-flash", api_key=api_key)
        Settings.embed_model = GeminiEmbedding(model_name="models/text-embedding-004", api_key=api_key)
        Settings.node_parser = SentenceSplitter(chunk_size=800, chunk_overlap=20)
        Settings.num_output = 512
        Settings.context_window = 3900
        
        # Create document and index
        document = Document(text=text)
        index = VectorStoreIndex.from_documents([document])
        
        logging.info("Index created successfully")
        return index
    except Exception as e:
        st.error(f"Error creating index: {str(e)}")
        logging.error(f"Index creation error: {str(e)}")
        return None

def query_index(index, question):
    """Query the index with a question"""
    try:
        query_engine = index.as_query_engine()
        response = query_engine.query(question)
        return str(response)
    except Exception as e:
        st.error(f"Error querying index: {str(e)}")
        logging.error(f"Query error: {str(e)}")
        return None

def simple_qa_system(document_text, question):
    """Fallback simple Q&A system using text matching"""
    if not document_text or not question:
        return "Please upload a document and ask a question."
    
    question_lower = question.lower()
    sentences = [s.strip() for s in document_text.split('.') if s.strip()]
    
    question_words = set(question_lower.split())
    relevant_sentences = []
    
    for sentence in sentences:
        sentence_words = set(sentence.lower().split())
        common_words = question_words.intersection(sentence_words)
        if len(common_words) > 0:
            relevance_score = len(common_words) / len(question_words)
            relevant_sentences.append((sentence, relevance_score))
    
    if relevant_sentences:
        relevant_sentences.sort(key=lambda x: x[1], reverse=True)
        top_sentences = [s[0] for s in relevant_sentences[:3]]
        return '. '.join(top_sentences) + '.'
    else:
        return "I couldn't find specific information related to your question in the document."

def main():
    st.set_page_config(page_title="QA System", page_icon="🤖", layout="wide")
    st.title("🤖 QA with Documents (Gemini-Powered)")
    
    # Sidebar for API Key and Info
    with st.sidebar:
        st.header("⚙️ Configuration")
        
        # Try to get API key from .env
        api_key_from_env = os.getenv("GOOGLE_API_KEY")
        
        if api_key_from_env:
            st.success("✅ API Key loaded from .env file")
            api_key = api_key_from_env
            use_manual = st.checkbox("Override with manual API key", value=False)
            if use_manual:
                api_key = st.text_input("Google API Key", type="password", value="")
        else:
            st.warning("⚠️ No API Key found in .env file")
            api_key = st.text_input("Google API Key", type="password", help="Enter your Google API key")
            
        st.markdown("---")
        st.header("ℹ️ About")
        st.write("This QA system uses Google's Gemini LLM for intelligent document Q&A.")
        
        st.header("🔧 Features")
        st.write("✓ PDF, TXT, DOCX support")
        st.write("✓ Gemini 2.5 Flash LLM")
        st.write("✓ Semantic search")
        st.write("✓ Document statistics")
        
        st.markdown("---")
        st.header("🔑 Get API Key")
        st.markdown("[Get Google API Key →](https://makersuite.google.com/app/apikey)")
        
        st.markdown("---")
        st.header("📂 Current Directory")
        st.code(os.getcwd())
    
    # Display status
    col1, col2 = st.columns(2)
    with col1:
        if LLAMA_INDEX_AVAILABLE:
            st.success("✅ LlamaIndex Available")
        else:
            st.error("❌ LlamaIndex Not Available")
    
    with col2:
        if api_key:
            st.success("✅ API Key Configured")
        else:
            st.warning("⚠️ No API Key")
    
    # Initialize session state
    if 'document_text' not in st.session_state:
        st.session_state.document_text = ""
    if 'index' not in st.session_state:
        st.session_state.index = None
    if 'file_processed' not in st.session_state:
        st.session_state.file_processed = False
    
    # File upload section
    st.header("📁 Upload Document")
    
    uploaded_file = st.file_uploader(
        "Choose a file",
        type=['pdf', 'txt', 'docx'],
        help="Supported formats: PDF, TXT, DOCX"
    )
    
    if uploaded_file is not None:
        # Check if this is a new file
        file_id = f"{uploaded_file.name}_{uploaded_file.size}"
        if 'last_file_id' not in st.session_state or st.session_state.last_file_id != file_id:
            st.session_state.last_file_id = file_id
            st.session_state.file_processed = False
            st.session_state.index = None
        
        if not st.session_state.file_processed:
            with st.spinner("📄 Processing document..."):
                document_text = load_data_from_upload(uploaded_file)
                st.session_state.document_text = document_text
                st.session_state.file_processed = True
            
            if document_text:
                st.success(f"✅ Document processed: {len(document_text):,} characters")
                
                # Create index if API key is provided
                if api_key and LLAMA_INDEX_AVAILABLE:
                    with st.spinner("🔄 Creating semantic index with Gemini..."):
                        index = create_index_from_text(document_text, api_key)
                        if index:
                            st.session_state.index = index
                            st.success("✅ Gemini index created!")
                        else:
                            st.warning("⚠️ Index creation failed. Using keyword matching.")
                elif not api_key:
                    st.info("💡 Add API key to enable Gemini-powered answers")
        
        # Show preview
        if st.session_state.document_text:
            with st.expander("📖 Document Preview"):
                preview_text = st.session_state.document_text[:1000]
                if len(st.session_state.document_text) > 1000:
                    preview_text += "..."
                st.text_area("First 1000 characters:", value=preview_text, height=200, disabled=True)
    
    # Q&A Section
    if st.session_state.document_text:
        st.markdown("---")
        st.header("❓ Ask Questions")
        
        # Show mode indicator
        if st.session_state.index and api_key:
            st.info("🤖 **Mode:** Gemini-powered intelligent answers")
        else:
            st.warning("⚠️ **Mode:** Basic keyword matching (add API key for better results)")
        
        # Question input
        question = st.text_input("💭 Your question:", placeholder="e.g., What is machine learning?")
        
        col1, col2 = st.columns([1, 5])
        with col1:
            ask_button = st.button("🔍 Get Answer", type="primary", use_container_width=True)
        with col2:
            clear_button = st.button("🗑️ Clear", use_container_width=True)
        
        if clear_button:
            st.rerun()
        
        if ask_button and question:
            with st.spinner("🤔 Thinking..."):
                # Use Gemini if available
                if st.session_state.index and api_key:
                    answer = query_index(st.session_state.index, question)
                    if not answer:
                        answer = simple_qa_system(st.session_state.document_text, question)
                else:
                    answer = simple_qa_system(st.session_state.document_text, question)
                
            st.markdown("### 📝 Answer:")
            st.info(answer)
        
        # Pre-defined questions
        st.markdown("---")
        st.subheader("💡 Try These Questions")
        
        col1, col2, col3 = st.columns(3)
        
        with col1:
            if st.button("📋 Summarize Document", use_container_width=True):
                with st.spinner("Generating summary..."):
                    if st.session_state.index and api_key:
                        answer = query_index(st.session_state.index, "Provide a brief summary of this document")
                    else:
                        answer = simple_qa_system(st.session_state.document_text, "main topic summary overview")
                st.markdown("**Summary:**")
                st.info(answer)
        
        with col2:
            if st.button("📊 Key Points", use_container_width=True):
                with st.spinner("Finding key points..."):
                    if st.session_state.index and api_key:
                        answer = query_index(st.session_state.index, "What are the key points and main findings?")
                    else:
                        answer = simple_qa_system(st.session_state.document_text, "key points important findings main")
                st.markdown("**Key Points:**")
                st.info(answer)
        
        with col3:
            if st.button("🎯 Main Topic", use_container_width=True):
                with st.spinner("Identifying topic..."):
                    if st.session_state.index and api_key:
                        answer = query_index(st.session_state.index, "What is the main topic of this document?")
                    else:
                        answer = simple_qa_system(st.session_state.document_text, "main topic subject about")
                st.markdown("**Main Topic:**")
                st.info(answer)
        
        # Document stats
        st.markdown("---")
        st.header("📊 Document Statistics")
        col1, col2, col3, col4 = st.columns(4)
        
        with col1:
            st.metric("📝 Characters", f"{len(st.session_state.document_text):,}")
        with col2:
            word_count = len(st.session_state.document_text.split())
            st.metric("📖 Words", f"{word_count:,}")
        with col3:
            sentences = len([s for s in st.session_state.document_text.split('.') if s.strip()])
            st.metric("💬 Sentences", f"{sentences:,}")
        with col4:
            lines = len(st.session_state.document_text.split('\n'))
            st.metric("📄 Lines", f"{lines:,}")
    
    else:
        # Welcome screen
        st.info("👆 Upload a document above to get started!")
        
        st.markdown("---")
        st.header("🚀 How to Use")
        
        col1, col2 = st.columns(2)
        
        with col1:
            st.markdown("""
            ### 📝 Step 1: Prepare
            - Ensure your `.env` file has `GOOGLE_API_KEY`
            - Or enter API key manually in sidebar
            - Get key from [Google AI Studio](https://makersuite.google.com/app/apikey)
            """)
            
            st.markdown("""
            ### 📤 Step 2: Upload
            - Click "Browse files" above
            - Select PDF, TXT, or DOCX file
            - Wait for processing
            """)
        
        with col2:
            st.markdown("""
            ### 💬 Step 3: Ask Questions
            - Type your question naturally
            - Click "Get Answer"
            - Or use quick question buttons
            """)
            
            st.markdown("""
            ### 🎯 Tips
            - Gemini mode gives better answers
            - Larger documents take longer
            - Try specific questions for best results
            """)

if __name__ == "__main__":
    main()